from docx import Document

doc = Document(r"public\KO'P BERILADIGAN SAVOLLAR.DOCX")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print("\n--- All paragraphs ---")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"{i}: {p.text}")